#!/usr/bin/env python3
"""
Generate grouped AI News Word document.

Groups articles into: OpenAI | Anthropic | BigTech | Other
BigTech = Google/DeepMind, Apple, Amazon/AWS, Meta, Microsoft, Netflix, xAI/Grok, NVIDIA

Output: one table with group header rows, followed by a fundraising table.
Usage:
  python tools/generate_ai_doc.py --start_date 2026-03-31 --end_date 2026-04-09 \\
    --articles .tmp/summarized_articles.json [--chinese-only | --translate] \\
    --output-prefix AI_News
"""

import os
import sys
import json
import re
import argparse
import time
from datetime import datetime
from dotenv import load_dotenv

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from tools.utils import format_date_for_display
from tools.generate_word_doc import (
    add_hyperlink,
    add_formatted_text,
    translate_to_chinese_claude,
    extract_funding_with_openai,
    create_funding_table,
    convert_bullets_to_paragraph,
)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import requests
except ImportError:
    print("ERROR: requests not installed. Run: pip install requests")
    sys.exit(1)


# ── Fundraising detection ──────────────────────────────────────────────────────

FUNDRAISING_KEYWORDS = [
    "raises", "raised", "raise", "funding", "fundrais",
    "series a", "series b", "series c", "series d", "series e",
    "seed round", "pre-seed", "venture", "investment round",
    "valuation", "unicorn", "acqui", "ipo", "goes public",
    "million", "billion", "$", "融资", "投资", "轮", "估值", "收购",
]


def is_fundraising_article(_article: dict) -> bool:
    """Disabled — all articles go to the main news table."""
    return False


def article_to_funding_event(article: dict, claude_key: str = None) -> dict:
    """Use Claude to extract structured funding info from a detected funding article."""
    title = article.get("title", "")
    # Prefer full article content over the condensed summary so founder details aren't lost
    body = article.get("content") or article.get("summary") or article.get("description", "")
    date = article.get("published_at", "")[:10]

    if claude_key:
        try:
            prompt = (
                f"从以下新闻文章中提取融资信息，所有字段均用中文填写。\n\n"
                f"标题：{title}\n文章正文：{body}\n\n"
                f"返回一个JSON对象，包含以下字段：\n"
                f'"company": 获得融资的公司名称\n'
                f'"summary": 用中文描述该公司，包含：(1) 一句话说明公司核心业务，(2) 如文章中提到创始人背景信息，请附上（姓名、曾任职的公司和职责）。'
                f'参考格式："AI-native 网络安全公司，用 AI agent 实时检测攻击并自动响应。创始人 XX 曾负责 Amazon Web Services GuardDuty"\n'
                f'"stage": 融资轮次（天使轮、Pre-A轮、A轮、B轮、C轮等，收购填"收购"，未知填"不详"）\n'
                f'"raise": 融资金额（例如："5000万美元"，未知填"不详"）\n'
                f'"valuation": 融资后估值（例如："5亿美元"，未知填"不详"）\n'
                f'"investors": 主要投资方（例如："领投：红杉资本"，未知填"不详"）\n\n'
                f"仅返回有效JSON，不要包含其他文字。"
            )
            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": claude_key,
                    "anthropic-version": "2023-06-01",
                },
                json={
                    "model": "claude-haiku-4-5-20251001",
                    "max_tokens": 300,
                    "messages": [{"role": "user", "content": prompt}],
                },
                timeout=30,
            )
            response.raise_for_status()
            text = response.json()["content"][0]["text"].strip()
            text = re.sub(r"^```json\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
            data = json.loads(text)
            return {
                "date": date,
                "company": data.get("company", title),
                "summary": data.get("summary", ""),
                "stage": data.get("stage", "N/A"),
                "raise": data.get("raise", "N/A"),
                "valuation": data.get("valuation", "N/A"),
                "investors": data.get("investors", "N/A"),
                "_url": article.get("url", ""),
            }
        except Exception as e:
            print(f"    WARNING: extraction failed ({e}), using fallback")

    # Fallback: no API key or extraction failed
    raw = body.split(".")[0] if body else title
    return {
        "date": date,
        "company": title,
        "summary": raw[:200],
        "stage": "N/A",
        "raise": "N/A",
        "valuation": "N/A",
        "investors": "N/A",
        "_url": article.get("url", ""),
    }


# ── Group definitions ──────────────────────────────────────────────────────────

GROUP_ORDER = ["OpenAI", "Anthropic", "BigTech", "Other"]

GROUP_DISPLAY = {
    "OpenAI": "OpenAI",
    "Anthropic": "Anthropic",
    "BigTech": "BigTech  ( Google · Apple · Amazon · Meta · Microsoft · Netflix · xAI · NVIDIA )",
    "Other": "Other",
}

# Header row fill colours (light blue family)
GROUP_COLORS = {
    "OpenAI":    "D6E4F0",
    "Anthropic": "D6E4F0",
    "BigTech":   "D6E4F0",
    "Other":     "D6E4F0",
}

OPENAI_KEYWORDS = [
    "openai", "chatgpt", "gpt-4", "gpt-3", "gpt4", "gpt3", " gpt ",
    "sora", " o1 ", " o1\n", " o3 ", " o3\n", " o4 ", " o4\n",
    "dall-e", "dall·e", "whisper", "altman", "sam altman",
]
ANTHROPIC_KEYWORDS = ["anthropic", "claude"]
BIGTECH_KEYWORDS = [
    # Google
    "google", "谷歌", "deepmind", "gemini", "bard", "waymo",
    # Apple
    "apple", "苹果",
    # Amazon
    "amazon", "亚马逊", " aws ",
    # Meta
    " meta ", "meta\n", "facebook", "脸书", "instagram", "whatsapp", "llama",
    # Microsoft
    "microsoft", "微软", "bing", " azure ", "copilot",
    # Netflix
    "netflix", "奈飞",
    # xAI / Elon
    " xai ", "x.ai", "grok", "elon musk", "马斯克",
    # NVIDIA
    "nvidia", "英伟达", "cuda",
]


def classify_article(article: dict) -> str:
    """Classify article into OpenAI | Anthropic | BigTech | Other by keyword match."""
    text = " " + " ".join([
        article.get("title", ""),
        article.get("summary", ""),
        article.get("description", ""),
    ]).lower() + " "

    for kw in OPENAI_KEYWORDS:
        if kw in text:
            return "OpenAI"
    for kw in ANTHROPIC_KEYWORDS:
        if kw in text:
            return "Anthropic"
    for kw in BIGTECH_KEYWORDS:
        if kw in text:
            return "BigTech"
    return "Other"


# ── Table helpers ──────────────────────────────────────────────────────────────

def _add_group_header_row(table, label: str, fill_hex: str = "D6E4F0"):
    """Add a full-width merged header row with coloured background."""
    row = table.add_row()
    # Merge both cells
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]

    # Background shading
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

    # Text
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 73, 125)


def create_grouped_news_table(
    doc: Document,
    articles: list,
    chinese_only: bool = False,
    translate: bool = False,
    claude_key: str = None,
):
    """
    Build the grouped AI News table.

    Groups: OpenAI → Anthropic → BigTech → Other
    Within each group, articles are sorted oldest first.
    """
    # Classify and bucket
    groups = {g: [] for g in GROUP_ORDER}
    for article in articles:
        groups[classify_article(article)].append(article)

    # Sort within each group by date
    for g in GROUP_ORDER:
        groups[g].sort(key=lambda x: x.get("published_at", ""))

    total = sum(len(v) for v in groups.values())

    heading = doc.add_heading("AI News Summary", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f"Total: {total} articles\n")

    # Create 2-column table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(6.0)

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Summary"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(12)

    # Populate groups
    article_counter = 0
    for group in GROUP_ORDER:
        group_articles = groups[group]
        if not group_articles:
            continue

        _add_group_header_row(table, GROUP_DISPLAY[group], GROUP_COLORS[group])

        for article in group_articles:
            article_counter += 1
            if translate:
                print(f"  [{article_counter}/{total}] Translating ({group})...")

            row_cells = table.add_row().cells

            # Col 0: date
            date_str = format_date_for_display(article.get("published_at", ""))
            date_run = row_cells[0].paragraphs[0].add_run(date_str)
            date_run.font.size = Pt(10)

            # Col 1: hyperlinked title + summary
            summary_cell = row_cells[1]
            summary_para = summary_cell.paragraphs[0]

            title = article.get("title", "No title")
            url = article.get("url", "")
            if url:
                add_hyperlink(summary_para, url, title)
            else:
                run = summary_para.add_run(title)
                run.bold = True
                run.font.size = Pt(10)

            raw_summary = article.get("summary", article.get("description", "No summary available"))
            summary_text = convert_bullets_to_paragraph(raw_summary)

            summary_para.add_run("\n\n")

            if chinese_only:
                add_formatted_text(summary_para, summary_text, font_size=10)
            elif translate and claude_key:
                chinese = translate_to_chinese_claude(claude_key, summary_text)
                if chinese:
                    add_formatted_text(summary_para, chinese, font_size=10)
                    time.sleep(0.3)
                summary_para.add_run("\n\n")
                add_formatted_text(summary_para, summary_text, font_size=10)
            else:
                add_formatted_text(summary_para, summary_text, font_size=10)

    return table


# ── Main generator ─────────────────────────────────────────────────────────────

def generate_ai_doc(
    start_date: str,
    end_date: str,
    articles_file: str = ".tmp/summarized_articles.json",
    output_dir: str = "output",
    max_articles: int = None,
    translate: bool = False,
    chinese_only: bool = False,
    output_prefix: str = "AI_News",
    no_funding: bool = False,
    funding_wechat_file: str = None,
):
    load_dotenv()

    print("Loading data...")
    if not os.path.exists(articles_file):
        print(f"ERROR: {articles_file} not found")
        sys.exit(1)

    with open(articles_file, "r", encoding="utf-8") as f:
        articles = json.load(f)

    if max_articles:
        articles = articles[:max_articles]

    print(f"Loaded {len(articles)} articles")

    # Separate fundraising articles from regular news (unless disabled)
    if no_funding:
        regular_articles = articles
        funding_articles = []
    else:
        regular_articles = [a for a in articles if not is_fundraising_article(a)]
        funding_articles = [a for a in articles if is_fundraising_article(a)]
        if funding_articles:
            print(f"  Detected {len(funding_articles)} fundraising articles — moving to funding table")

    claude_key = None
    if translate:
        claude_key = os.getenv("ANTHROPIC_API_KEY")
        if not claude_key:
            print("ERROR: ANTHROPIC_API_KEY not found in .env")
            sys.exit(1)
        print("Translation enabled (Claude)")

    openai_key = os.getenv("OPENAI_API_KEY")

    print("Creating Word document...")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    title_para = doc.add_heading("AI News Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"{start_date}  to  {end_date}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    total_str = str(len(regular_articles))
    if not no_funding and funding_articles:
        total_str += f" (+ {len(funding_articles)} moved to funding table)"
    doc.add_paragraph(f"Total Articles: {total_str}")
    doc.add_paragraph("")

    print("Creating grouped news table...")
    create_grouped_news_table(doc, regular_articles, chinese_only, translate, claude_key)

    if not no_funding:
        anthropic_key = claude_key or os.getenv("ANTHROPIC_API_KEY")

        # Load WeChat fundraising articles if provided
        wechat_funding = []
        if funding_wechat_file and os.path.exists(funding_wechat_file):
            with open(funding_wechat_file, "r", encoding="utf-8") as f:
                wechat_funding_articles = json.load(f)
            print(f"Extracting funding details from {len(wechat_funding_articles)} WeChat fundraising article(s)...")
            for i, a in enumerate(wechat_funding_articles, 1):
                print(f"  [{i}/{len(wechat_funding_articles)}] {a.get('title', '')[:70]}...")
                wechat_funding.append(article_to_funding_event(a, anthropic_key))

        print("Searching for AI funding news with ChatGPT (day by day)...")
        if openai_key:
            funding_events = extract_funding_with_openai(openai_key, start_date, end_date)
            all_funding = funding_events + wechat_funding
            print(f"  Found {len(funding_events)} from web search + {len(wechat_funding)} from WeChat")
            create_funding_table(doc, all_funding)
        elif wechat_funding:
            print("  OPENAI_API_KEY not set — using WeChat fundraising articles only")
            create_funding_table(doc, wechat_funding)
        else:
            print("  WARNING: OPENAI_API_KEY not set — skipping funding section")

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_prefix}_{start_date.replace('-','')}_{end_date.replace('-','')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    print(f"\n✓ Document saved to {filepath}")
    return filepath


def main():
    parser = argparse.ArgumentParser(description="Generate grouped AI News Word document")
    parser.add_argument("--start_date", required=True, help="Start date (YYYY-MM-DD)")
    parser.add_argument("--end_date", required=True, help="End date (YYYY-MM-DD)")
    parser.add_argument("--articles", default=".tmp/summarized_articles.json")
    parser.add_argument("--output_dir", default="output")
    parser.add_argument("--max", type=int, default=None)
    parser.add_argument("--translate", action="store_true", help="Add Chinese translation")
    parser.add_argument("--chinese-only", action="store_true", help="Chinese summary only")
    parser.add_argument("--no-funding", action="store_true", help="Skip funding detection and funding table")
    parser.add_argument("--output-prefix", default="AI_News")
    parser.add_argument("--funding-wechat", default=None, help="Path to summarized WeChat fundraising articles JSON")
    args = parser.parse_args()

    generate_ai_doc(
        args.start_date,
        args.end_date,
        args.articles,
        args.output_dir,
        args.max,
        args.translate,
        args.chinese_only,
        args.output_prefix,
        args.no_funding,
        args.funding_wechat,
    )


if __name__ == "__main__":
    main()
