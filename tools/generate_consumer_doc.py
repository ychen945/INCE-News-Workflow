#!/usr/bin/env python3
"""
Generate Word document for Consumer Tech News.

Format:
  Section 1 — 行业动态 (Industry Updates)
    Article Title (hyperlinked, bold heading)
    1-2 paragraph Chinese summary

  Section 2 — 融资新闻 (Funding News)
    Article Title (hyperlinked, bold heading)
    1-2 paragraph Chinese summary (includes product info, founders, revenue, size)

Articles must have a 'category' field set by summarize_articles.py --consumer mode.
No tables, no dates.

Output: output/Consumer_News_YYYYMMDD_YYYYMMDD.docx
"""

import os
import sys
import json
import re
import argparse
from datetime import datetime
from dotenv import load_dotenv

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def add_hyperlink(paragraph, url: str, text: str):
    """Add a clickable hyperlink to a paragraph (blue, underlined)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    b = OxmlElement('w:b')
    rPr.append(b)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt
    rPr.append(sz)

    run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def convert_bullets_to_paragraph(text: str) -> str:
    """Convert bullet point text to flowing paragraph format."""
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        line = line.strip()
        line = re.sub(r'^[\-\*•]\s*', '', line)
        line = re.sub(r'^\d+[\.\)]\s*', '', line)
        if line:
            cleaned.append(line)
    return ' '.join(cleaned)


def add_article(doc: Document, article: dict):
    """
    Add a single article to the document:
      - Bold hyperlinked title as a heading
      - Summary as paragraph text below
    """
    title = article.get('title', '（无标题）')
    url = article.get('url', '')
    summary = article.get('summary', article.get('description', '暂无摘要'))

    # Convert bullets to paragraph if needed
    summary_text = convert_bullets_to_paragraph(summary)

    # Title paragraph (acts as sub-heading)
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(10)
    title_para.paragraph_format.space_after = Pt(2)
    if url:
        add_hyperlink(title_para, url, title)
    else:
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(12)

    # Summary paragraph
    summary_para = doc.add_paragraph(summary_text)
    summary_para.paragraph_format.space_before = Pt(0)
    summary_para.paragraph_format.space_after = Pt(8)
    for run in summary_para.runs:
        run.font.size = Pt(10.5)


def generate_consumer_doc(start_date: str, end_date: str,
                           articles_file: str = '.tmp/summarized_wechat.json',
                           output_dir: str = 'output',
                           max_articles: int = None):
    """
    Main document generation function.

    Args:
        start_date: YYYY-MM-DD (used for filename and document subtitle)
        end_date: YYYY-MM-DD
        articles_file: Path to JSON with 'summary' and 'category' fields
        output_dir: Output directory
        max_articles: Cap on total articles (None = all)
    """
    load_dotenv()

    print("Loading data...")

    if not os.path.exists(articles_file):
        print(f"ERROR: {articles_file} not found")
        print("Run: python tools/summarize_articles.py --consumer --provider claude --yes")
        sys.exit(1)

    with open(articles_file, 'r', encoding='utf-8') as f:
        articles = json.load(f)

    if max_articles:
        articles = articles[:max_articles]

    print(f"Loaded {len(articles)} articles")

    # Split by category
    industry = [a for a in articles if a.get('category', '行业动态') == '行业动态']
    funding = [a for a in articles if a.get('category') == '融资新闻']

    print(f"  行业动态: {len(industry)} articles")
    print(f"  融资新闻: {len(funding)} articles")

    # Build document
    print("Creating Word document...")
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    # Document title
    title_para = doc.add_heading('消费科技新闻报告', level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f'{start_date} 至 {end_date}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('')

    # Section 1: 行业动态
    h1 = doc.add_heading('行业动态', level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if industry:
        for article in industry:
            add_article(doc, article)
    else:
        doc.add_paragraph('本期暂无行业动态。')

    doc.add_paragraph('')

    # Section 2: 融资新闻
    h2 = doc.add_heading('融资新闻', level=1)
    h2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if funding:
        for article in funding:
            add_article(doc, article)
    else:
        doc.add_paragraph('本期暂无融资新闻。')

    # Save
    os.makedirs(output_dir, exist_ok=True)
    filename = f'Consumer_News_{start_date.replace("-", "")}_{end_date.replace("-", "")}.docx'
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    print(f"\n✓ Document saved to {filepath}")
    return filepath


def main():
    parser = argparse.ArgumentParser(description='Generate Consumer Tech News Word document')
    parser.add_argument('--start_date', required=True, help='Start date (YYYY-MM-DD)')
    parser.add_argument('--end_date', required=True, help='End date (YYYY-MM-DD)')
    parser.add_argument('--articles', default='.tmp/summarized_wechat.json',
                        help='Summarized articles file (default: .tmp/summarized_wechat.json)')
    parser.add_argument('--output_dir', default='output', help='Output directory (default: output)')
    parser.add_argument('--max', type=int, default=None, help='Maximum articles to include (default: all)')
    args = parser.parse_args()

    generate_consumer_doc(args.start_date, args.end_date, args.articles, args.output_dir, args.max)


if __name__ == '__main__':
    main()
