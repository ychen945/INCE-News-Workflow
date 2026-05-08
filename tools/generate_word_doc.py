#!/usr/bin/env python3
"""
Generate Word document with AI News table:
- Title (hyperlinked) + Date + Source | Summary (with optional Chinese translation)

Uses python-docx for formatting
"""

import os
import sys
import json
import re
import argparse
import time
from datetime import datetime
from dotenv import load_dotenv

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from tools.utils import format_date_for_display

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import requests
except ImportError:
    print("ERROR: requests not installed. Run: pip install requests")
    sys.exit(1)


def add_hyperlink(paragraph, url: str, text: str):
    """
    Add a hyperlink to a paragraph

    python-docx doesn't have native hyperlink support, so we use XML

    Args:
        paragraph: docx paragraph object
        url: URL to link to
        text: Link text
    """
    # Create hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run element
    run = OxmlElement('w:r')

    # Run properties (style)
    rPr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    run.append(rPr)

    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_formatted_text(paragraph, text: str, font_size: int = 10):
    """
    Add text to paragraph with markdown bold (**text**) converted to Word bold

    Args:
        paragraph: docx paragraph object
        text: Text that may contain **bold** markdown
        font_size: Font size in points
    """
    # Pattern to match **bold** text
    pattern = r'\*\*(.+?)\*\*'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the bold part
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.size = Pt(font_size)

        # Add bold text
        run = paragraph.add_run(match.group(1))
        run.bold = True
        run.font.size = Pt(font_size)

        last_end = match.end()

    # Add remaining text after last match
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(font_size)


def translate_to_chinese_claude(api_key: str, text: str) -> str:
    """
    Translate text to Chinese using Claude API

    Args:
        api_key: Anthropic API key
        text: Text to translate

    Returns:
        Chinese translation
    """
    try:
        url = "https://api.anthropic.com/v1/messages"

        headers = {
            'Content-Type': 'application/json',
            'x-api-key': api_key,
            'anthropic-version': '2023-06-01'
        }

        payload = {
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 500,
            "messages": [
                {
                    "role": "user",
                    "content": f"Translate the following text to Simplified Chinese. Only output the translation, nothing else.\n\n{text}"
                }
            ]
        }

        response = requests.post(url, json=payload, headers=headers, timeout=30)
        response.raise_for_status()

        result = response.json()

        if 'content' in result and len(result['content']) > 0:
            return result['content'][0]['text'].strip()
        else:
            print(f"  WARNING: Unexpected Claude response format")
            return ""

    except Exception as e:
        print(f"  WARNING: Translation failed: {e}")
        return ""


def _search_funding_single_day(api_key: str, date: str, topic: str) -> list:
    """
    Search for funding events on a single date using OpenAI web search.
    Returns list of funding event dicts.
    """
    if topic == 'deeptech':
        sector_desc = "深科技公司（包括机器人、先进材料、量子计算、生物/医疗科技、航天、半导体、清洁能源等硬科技领域）"
    else:
        sector_desc = "AI / 人工智能公司"

    prompt = f"""搜索网络，找出{date}宣布的{sector_desc}融资轮次、投资和收购事件。

对于每个融资事件，返回一个JSON对象。返回一个包含以下字段的JSON数组：
- "date": 宣布日期，格式为YYYY-MM-DD
- "company": 获得融资的公司名称
- "summary": 用中文描述该公司，包含：(1) 一句话说明公司的核心业务，(2) 如网上有创始人相关背景信息，请附上（例如：曾就职的知名公司、负责的项目、相关行业经验等）。参考格式："AI-native 网络安全公司，用 AI agent 实时检测攻击并自动响应。创始人 XX 曾负责 Amazon Web Services GuardDuty，联合创始人 YY 曾在 Abnormal AI 负责机器学习"
- "stage": 融资轮次（天使轮、Pre-A轮、A轮、B轮、C轮等，如为收购则填"收购"，未知填"不详"）
- "raise": 融资金额（例如："5000万美元"、"12亿美元"，未知填"不详"）
- "valuation": 融资后估值（例如："5亿美元"、"12亿美元"，未知填"不详"）
- "investors": 主要投资方（例如："领投：红杉资本，跟投：Andreessen Horowitz"，未知填"不详"）
- "url": 最相关的新闻来源链接（如有则填完整URL，否则填""）

只包含实际融资事件（已筹集资金、收购、IPO）。如未找到任何事件，返回空数组[]。
仅返回有效的JSON数组，不要包含其他文字。"""

    try:
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {api_key}'
        }
        payload = {
            "model": "gpt-4o-search-preview",
            "web_search_options": {},
            "messages": [{"role": "user", "content": prompt}],
        }

        for attempt in range(3):
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                json=payload, headers=headers, timeout=90
            )
            if response.status_code == 429:
                wait = 20 * (attempt + 1)
                print(f"  Rate limited, waiting {wait}s...")
                time.sleep(wait)
                continue
            break

        result = response.json()
        if 'choices' not in result:
            print(f"  OpenAI response: {result}")
            return []

        content = result['choices'][0]['message']['content'].strip()

        json_match = re.search(r'```(?:json)?\s*(\[.*?\])\s*```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)
        else:
            array_match = re.search(r'(\[.*\])', content, re.DOTALL)
            if array_match:
                content = array_match.group(1)
            else:
                return []

        events = json.loads(content)
        return events if isinstance(events, list) else []

    except Exception as e:
        print(f"  WARNING: Funding search failed for {date}: {e}")
        return []


def extract_funding_with_openai(api_key: str, start_date: str, end_date: str,
                               topic: str = 'ai') -> list:
    """
    Use OpenAI with web search to find funding events in a date range.
    Searches day-by-day to avoid the model skipping dates in long ranges.

    Args:
        api_key: OpenAI API key
        start_date: YYYY-MM-DD start date
        end_date: YYYY-MM-DD end date
        topic: 'ai' (default) or 'deeptech'

    Returns:
        List of funding event dicts with keys: date, company, summary, stage, raise, valuation, investors
    """
    from datetime import timedelta

    start_dt = datetime.strptime(start_date, "%Y-%m-%d")
    end_dt = datetime.strptime(end_date, "%Y-%m-%d")

    total_days = (end_dt - start_dt).days + 1
    print(f"  Searching {total_days} day(s): {start_date} to {end_date}")

    all_events = []
    current = start_dt
    day_num = 0
    while current <= end_dt:
        day_num += 1
        date_str = current.strftime("%Y-%m-%d")
        print(f"  [{day_num}/{total_days}] Searching {date_str}...")
        events = _search_funding_single_day(api_key, date_str, topic)
        print(f"    Found {len(events)} events")
        all_events.extend(events)
        current += timedelta(days=1)
        time.sleep(1.0)  # rate limiting between days

    # Deduplicate by company + date
    seen = set()
    unique = []
    for e in all_events:
        key = (e.get('company', '').lower().strip(), e.get('date', '')[:10])
        if key not in seen:
            seen.add(key)
            unique.append(e)

    print(f"  Total unique funding events: {len(unique)}")
    return unique


def create_funding_table(doc: Document, funding_events: list, heading: str = 'AI 融资动态'):
    """
    Add Fundraising News section with a 7-column table (all Chinese).

    Args:
        doc: Document object
        funding_events: List of funding event dicts
        heading: Section heading text
    """
    doc.add_paragraph('')
    h = doc.add_heading(heading, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not funding_events:
        doc.add_paragraph('该时间段内未发现融资事件。')
        return

    doc.add_paragraph(f'共 {len(funding_events)} 条融资记录\n')

    # Create table with 7 columns
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'

    # Set column widths
    col_widths = [Inches(0.85), Inches(1.0), Inches(2.0), Inches(0.65), Inches(0.65), Inches(0.75), Inches(1.3)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Header row (Chinese)
    headers = ['日期', '公司', '概述', '轮次', '融资额', '估值', '投资方']
    header_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        header_cells[i].text = h_text
        for run in header_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    # Sort by date oldest first
    funding_events.sort(key=lambda x: x.get('date', ''))

    # Data rows
    for event in funding_events:
        row_cells = table.add_row().cells

        # Col 0: date
        date_run = row_cells[0].paragraphs[0].add_run(event.get('date', '')[:10])
        date_run.font.size = Pt(9)

        # Col 1: company name, hyperlinked if URL available
        company = event.get('company', '不详')
        url = event.get('url', event.get('_url', ''))
        company_para = row_cells[1].paragraphs[0]
        if url:
            add_hyperlink(company_para, url, company)
            for r in company_para.runs:
                r.font.size = Pt(9)
        else:
            run = company_para.add_run(company)
            run.font.size = Pt(9)

        # Cols 2-6: remaining fields
        remaining = [
            event.get('summary', '不详'),
            event.get('stage', '不详'),
            event.get('raise', '不详'),
            event.get('valuation', '不详'),
            event.get('investors', '不详'),
        ]
        for i, val in enumerate(remaining, start=2):
            para = row_cells[i].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)


def convert_bullets_to_paragraph(text: str) -> str:
    """
    Convert bullet point text to paragraph format.
    Removes bullet markers and joins into flowing text.

    Args:
        text: Text with potential bullet points

    Returns:
        Text as paragraph without bullets
    """
    # Remove common bullet markers and clean up
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        # Remove bullet markers
        line = re.sub(r'^[\-\*•]\s*', '', line)
        line = re.sub(r'^\d+[\.\)]\s*', '', line)
        if line:
            cleaned_lines.append(line)

    # Join with spaces to form paragraph
    return ' '.join(cleaned_lines)


# ── Deeptech category grouping ─────────────────────────────────────────────────

DEEPTECH_CATEGORY_ORDER = ["半导体", "机器人", "新能源", "其他"]

DEEPTECH_CATEGORY_COLORS = {
    "半导体": "E8F0FE",
    "机器人": "E6F4EA",
    "新能源": "FFF8E1",
    "其他":   "F3E5F5",
}

SEMICONDUCTOR_KEYWORDS = [
    "semiconductor", "chip", "芯片", "半导体", "wafer", "fab", "foundry",
    "transistor", "lithography", "eda", "photonic", "asic", "fpga",
    "nvidia", "intel", "amd", "tsmc", "arm ", "risc", "memory", "dram",
    "nand", "soc", "gpu", "cpu", "mpu", "integrated circuit",
]
ROBOTICS_KEYWORDS = [
    "robot", "机器人", "humanoid", "autonomous", "drone", "无人机",
    "unmanned", "exoskeleton", "cobots", "manipulation", "locomotion",
    "actuator", "servo", "mechatronics", "automation",
]
ENERGY_KEYWORDS = [
    "energy", "新能源", "electric vehicle", "ev ", " ev\n", "solar",
    "battery", "电池", "储能", "充电", "wind", "nuclear", "hydrogen",
    "fuel cell", "grid", "power", "renewable", "carbon", "climate",
    "clean tech", "cleantech", "charging", "inverter", "photovoltaic",
]


def classify_deeptech_article(article: dict) -> str:
    """Classify a deeptech article into 半导体 | 机器人 | 新能源 | 其他."""
    text = " " + " ".join([
        article.get("title", ""),
        article.get("summary", ""),
        article.get("description", ""),
    ]).lower() + " "

    for kw in SEMICONDUCTOR_KEYWORDS:
        if kw in text:
            return "半导体"
    for kw in ROBOTICS_KEYWORDS:
        if kw in text:
            return "机器人"
    for kw in ENERGY_KEYWORDS:
        if kw in text:
            return "新能源"
    return "其他"


def _add_deeptech_header_row(table, label: str, fill_hex: str):
    """Add a full-width merged header row for deeptech category sections."""
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 73, 125)


def create_grouped_deeptech_table(
    doc: Document,
    articles: list,
    chinese_only: bool = False,
    translate: bool = False,
    claude_key: str = None,
    heading: str = '深科技新闻摘要',
):
    """Build a grouped deeptech news table: 半导体 → 机器人 → 新能源 → 其他."""
    groups = {cat: [] for cat in DEEPTECH_CATEGORY_ORDER}
    for article in articles:
        groups[classify_deeptech_article(article)].append(article)

    for cat in DEEPTECH_CATEGORY_ORDER:
        groups[cat].sort(key=lambda x: x.get("published_at", ""))

    total = sum(len(v) for v in groups.values())

    h = doc.add_heading(heading, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f'Total: {total} articles\n')

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(6.0)

    hdr = table.rows[0].cells
    hdr[0].text = 'Date'
    hdr[1].text = 'Summary'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(12)

    for cat in DEEPTECH_CATEGORY_ORDER:
        cat_articles = groups[cat]
        if not cat_articles:
            continue

        _add_deeptech_header_row(table, cat, DEEPTECH_CATEGORY_COLORS[cat])

        for article in cat_articles:
            row_cells = table.add_row().cells

            date_str = format_date_for_display(article.get('published_at', ''))
            row_cells[0].paragraphs[0].add_run(date_str).font.size = Pt(10)

            summary_para = row_cells[1].paragraphs[0]
            title = article.get('title', 'No title')
            url = article.get('url', '')
            if url:
                add_hyperlink(summary_para, url, title)
            else:
                run = summary_para.add_run(title)
                run.bold = True
                run.font.size = Pt(10)

            summary = article.get('summary', article.get('description', ''))
            summary_text = convert_bullets_to_paragraph(summary)
            summary_para.add_run('\n\n')

            if chinese_only:
                add_formatted_text(summary_para, summary_text, font_size=10)
            elif translate and claude_key:
                chinese = translate_to_chinese_claude(claude_key, summary_text)
                if chinese:
                    add_formatted_text(summary_para, chinese, font_size=10)
                    time.sleep(0.3)
                summary_para.add_run('\n\n')
                add_formatted_text(summary_para, summary_text, font_size=10)
            else:
                add_formatted_text(summary_para, summary_text, font_size=10)

    return table


def create_news_table(doc: Document, articles: list, max_articles: int = None, translate: bool = False, claude_key: str = None, chinese_only: bool = False):
    """
    Create AI News table with 2 columns: Date | Title + Summary

    Args:
        doc: Document object
        articles: List of articles with summaries
        max_articles: Maximum number of articles to include (None = all)
        translate: Whether to add Chinese translation
        claude_key: Anthropic API key for translation
    """
    # Sort by date (oldest first)
    articles.sort(key=lambda x: x.get('published_at', ''), reverse=False)

    # Limit articles if specified
    if max_articles:
        articles = articles[:max_articles]

    # Add heading
    heading = doc.add_heading('AI News Summary', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add count
    doc.add_paragraph(f'Total: {len(articles)} articles\n')

    # Create table with 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    # Set column widths
    table.columns[0].width = Inches(1.0)  # Date
    table.columns[1].width = Inches(6.0)  # Title + Summary

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Date'
    header_cells[1].text = 'Summary'

    # Format header
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)

    # Add data rows
    for i, article in enumerate(articles):
        if translate:
            print(f"  [{i+1}/{len(articles)}] Translating...")

        row_cells = table.add_row().cells

        # Column 1: Date only
        date_str = format_date_for_display(article.get('published_at', ''))
        date_cell = row_cells[0]
        date_para = date_cell.paragraphs[0]
        date_run = date_para.add_run(date_str)
        date_run.font.size = Pt(10)

        # Column 2: Title (hyperlinked) + Summary
        summary_cell = row_cells[1]
        summary_para = summary_cell.paragraphs[0]

        # Add hyperlinked title
        title = article.get('title', 'No title')
        url = article.get('url', '')
        if url:
            add_hyperlink(summary_para, url, title)
        else:
            run = summary_para.add_run(title)
            run.bold = True
            run.font.size = Pt(10)

        # Get summary and convert bullets to paragraph
        summary = article.get('summary', article.get('description', 'No summary available'))
        summary_paragraph = convert_bullets_to_paragraph(summary)

        if chinese_only:
            # Use summary as-is (already in Chinese from --language zh summarization)
            summary_para.add_run("\n\n")
            add_formatted_text(summary_para, summary_paragraph, font_size=10)
        elif translate and claude_key:
            # Add Chinese translation first, then English
            chinese = translate_to_chinese_claude(claude_key, summary_paragraph)
            if chinese:
                summary_para.add_run("\n\n")
                add_formatted_text(summary_para, chinese, font_size=10)
                time.sleep(0.5)  # Rate limiting for Claude
            summary_para.add_run("\n\n")
            add_formatted_text(summary_para, summary_paragraph, font_size=10)
        else:
            # English only
            summary_para.add_run("\n\n")
            add_formatted_text(summary_para, summary_paragraph, font_size=10)

    return table


def generate_word_doc(start_date: str, end_date: str,
                      articles_file: str = '.tmp/summarized_articles.json',
                      output_dir: str = 'output',
                      max_articles: int = None,
                      translate: bool = False,
                      chinese_only: bool = False,
                      output_prefix: str = 'AI_News',
                      funding_topic: str = 'ai',
                      doc_title: str = None):
    """
    Main document generation function

    Args:
        start_date: YYYY-MM-DD
        end_date: YYYY-MM-DD
        articles_file: Path to summarized articles JSON
        output_dir: Output directory
        max_articles: Maximum articles to include (None = all)
        translate: Add Chinese translation using Claude
    """
    load_dotenv()

    print("Loading data...")

    # Load summarized articles
    if not os.path.exists(articles_file):
        print(f"ERROR: {articles_file} not found")
        print("Make sure to run summarize_articles.py first")
        sys.exit(1)

    with open(articles_file, 'r', encoding='utf-8') as f:
        articles = json.load(f)

    print(f"Loaded {len(articles)} articles")

    # Check for Claude API key if translation is requested
    claude_key = None
    if translate:
        claude_key = os.getenv('ANTHROPIC_API_KEY')
        if not claude_key:
            print("ERROR: ANTHROPIC_API_KEY not found in .env file")
            sys.exit(1)
        print("Translation enabled (Claude)")

    # Check for OpenAI key for funding extraction
    openai_key = os.getenv('OPENAI_API_KEY')

    # Create document
    print("Creating Word document...")
    doc = Document()

    # Title
    display_title = doc_title if doc_title else 'AI News Report'
    title = doc.add_heading(display_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with date range
    subtitle = doc.add_paragraph(f'{start_date} to {end_date}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Metadata
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'Total Articles Collected: {len(articles)}')
    doc.add_paragraph('')

    # Create news table (grouped for deeptech, flat for others)
    print("Creating news table...")
    if funding_topic == 'deeptech':
        create_grouped_deeptech_table(doc, articles, chinese_only, translate, claude_key)
    else:
        create_news_table(doc, articles, max_articles, translate, claude_key, chinese_only)

    # Create funding section
    topic_label = "Deeptech" if funding_topic == "deeptech" else "AI"
    print(f"Searching for {topic_label} funding news with ChatGPT (live web search)...")
    if openai_key:
        funding_events = extract_funding_with_openai(openai_key, start_date, end_date, funding_topic)
        print(f"  Found {len(funding_events)} funding events")
        heading_map = {"AI": "AI 融资动态", "Deeptech": "深科技融资动态"}
        create_funding_table(doc, funding_events, heading=heading_map.get(topic_label, f"{topic_label} 融资动态"))
    else:
        print("  WARNING: OPENAI_API_KEY not set, skipping funding section")

    # Save document
    os.makedirs(output_dir, exist_ok=True)
    filename = f'{output_prefix}_{start_date.replace("-", "")}_{end_date.replace("-", "")}.docx'
    filepath = os.path.join(output_dir, filename)

    doc.save(filepath)

    print(f"\n✓ Document saved to {filepath}")
    return filepath


def main():
    """Main execution function"""
    parser = argparse.ArgumentParser(description='Generate AI News Word document')
    parser.add_argument('--start_date', required=True, help='Start date (YYYY-MM-DD)')
    parser.add_argument('--end_date', required=True, help='End date (YYYY-MM-DD)')
    parser.add_argument('--articles', default='.tmp/summarized_articles.json', help='Summarized articles file')
    parser.add_argument('--output_dir', default='output', help='Output directory')
    parser.add_argument('--max', type=int, default=None, help='Maximum articles to include (default: all)')
    parser.add_argument('--translate', action='store_true', help='Add Chinese translation using ChatGPT')
    parser.add_argument('--chinese-only', action='store_true', help='Output Chinese summary only (no English), for pre-summarized Chinese articles')
    parser.add_argument('--output-prefix', default='AI_News', help='Output filename prefix (default: AI_News)')
    parser.add_argument('--funding-topic', choices=['ai', 'deeptech'], default='ai',
                        help='Funding search topic: ai (default) or deeptech')
    parser.add_argument('--doc-title', default=None, help='Document title (default: AI News Report)')
    args = parser.parse_args()

    generate_word_doc(
        args.start_date,
        args.end_date,
        args.articles,
        args.output_dir,
        args.max,
        args.translate,
        args.chinese_only,
        args.output_prefix,
        args.funding_topic,
        args.doc_title,
    )


if __name__ == "__main__":
    main()
